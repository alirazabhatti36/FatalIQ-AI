from flask import Flask, request, send_file, send_from_directory, render_template, jsonify, redirect, abort, make_response
from PIL import Image
import pytesseract
import os, re, io, zipfile, shutil
import json
import tempfile
import importlib
import time
import csv
import html

# --- LinkedIn Optimizer import HATAYA ---
# from linkedin_optimizer import analyze_linkedin_profile, fetch_profile_from_url

# Tesseract path (Sirf Image OCR ke liye, agar Render pe nahi hai toh error dega)
WINDOWS_TESSERACT_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if os.name == 'nt' and os.path.exists(WINDOWS_TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = WINDOWS_TESSERACT_PATH

app = Flask(__name__)
app.url_map.strict_slashes = False
UPLOAD = 'uploads'
os.makedirs(UPLOAD, exist_ok=True)
UPLOAD_RETENTION_SECONDS = int(os.getenv('UPLOAD_RETENTION_SECONDS', '900'))

def purge_uploads(force=False, max_age_seconds=UPLOAD_RETENTION_SECONDS):
    now = time.time()
    try:
        for name in os.listdir(UPLOAD):
            path = os.path.join(UPLOAD, name)
            if not os.path.isfile(path):
                continue
            if force:
                os.remove(path)
                continue
            age = now - os.path.getmtime(path)
            if age > max_age_seconds:
                os.remove(path)
    except Exception:
        # Never block requests due to cleanup failures.
        pass

# Clear stale files on startup to avoid persistent leftovers across deploys.
purge_uploads(force=True)

@app.before_request
def cleanup_user_files():
    purge_uploads(force=False)

@app.after_request
def add_cors_headers(response):
    origin = request.headers.get('Origin')
    if origin in {'https://auratoolkit360.com', 'https://www.auratoolkit360.com'}:
        response.headers['Access-Control-Allow-Origin'] = origin
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        response.headers['Vary'] = 'Origin'

    # Privacy-first policy: prevent browser/proxy caching of user-generated files.
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

PAGE_TEMPLATES = {
    'index': 'index.html',
    'converter': 'converter.html',
    'resume-ats': 'resume-ats.html',
    'hr-helper': 'hr-helper.html',
    'privacy-policy': 'privacy-policy.html',
    'terms': 'terms.html',
    'contact': 'contact.html',
    'about': 'about.html',
    'blogs': 'blogs.html',
    'disclaimer': 'disclaimer.html',
}

BLOG_POST_TEMPLATES = {
    'ats-resume-keywords-guide-2026': 'blog/ats-resume-keywords-guide-2026.html',
    'resume-summary-examples-for-freshers': 'blog/resume-summary-examples-for-freshers.html',
    'linkedin-headline-formula-with-examples': 'blog/linkedin-headline-formula-with-examples.html',
    'pdf-to-word-formatting-fixes': 'blog/pdf-to-word-formatting-fixes.html',
    'hr-shortlisting-checklist-small-teams': 'blog/hr-shortlisting-checklist-small-teams.html',
    'privacy-first-file-conversion-guide': 'blog/privacy-first-file-conversion-guide.html',
}

def save(file, name):
    path = os.path.join(UPLOAD, name)
    file.save(path)
    return path

def is_windows_runtime():
    return os.name == 'nt'

def windows_only_response(feature_name):
    return jsonify({
        'error': f'{feature_name} is available only on a Windows server. For Render/Linux deployment, use the PDF, OCR, ATS, and browser-safe converters instead.'
    }), 501

def extract_text_from_docx(path):
    docx = importlib.import_module('docx')
    d = docx.Document(path)
    return ' '.join([p.text for p in d.paragraphs])

def extract_text_from_pdf(path):
    PyPDF2 = importlib.import_module('PyPDF2')
    text = ''
    with open(path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ''
    return text

def normalize_ats_text(text):
    stop_words = {
        'the','and','for','are','with','this','that','you','your','have','from',
        'will','can','was','they','their','been','has','but','not','all','our',
        'its','also','more','when','which','a','an','in','of','to','is','on','at',
        'we','us','he','she','it','them','his','her','these','those','as','by','be',
        'being','or','if','into','than','then','there','here','over','under','up',
        'down','out','per','via','about','after','before','during','between','within',
        'without','one','two','three','four','five','six','seven','eight','nine','ten'
    }
    words = re.findall(r'\b[a-z0-9+#.-]{2,}\b', (text or '').lower())
    return [word for word in words if word not in stop_words and len(word) > 2]

def extract_text_from_image(path):
    try:
        try:
            easyocr = importlib.import_module('easyocr')
            reader = easyocr.Reader(['en'])
            result = reader.readtext(path, detail=0)
            if result:
                return ' '.join(result)
        except Exception:
            pass

        return pytesseract.image_to_string(Image.open(path))
    except Exception:
        return ''

def extract_text_from_supported_file(path, ext=None):
    ext = (ext or os.path.splitext(path)[1].lstrip('.')).lower()
    if ext == 'pdf':
        return extract_text_from_pdf(path)
    if ext == 'docx':
        return extract_text_from_docx(path)
    if ext in {'txt', 'csv', 'md', 'rtf'}:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    if ext in {'png', 'jpg', 'jpeg', 'bmp', 'webp', 'tif', 'tiff'}:
        return extract_text_from_image(path)
    if ext == 'zip':
        combined_text = []
        with tempfile.TemporaryDirectory() as temp_dir:
            with zipfile.ZipFile(path, 'r') as archive:
                archive.extractall(temp_dir)
            for root, _, files in os.walk(temp_dir):
                for name in files:
                    inner_path = os.path.join(root, name)
                    inner_ext = os.path.splitext(name)[1].lstrip('.').lower()
                    if inner_ext in {'pdf','docx','txt','csv','md','rtf','png','jpg','jpeg','bmp','webp','tif','tiff'}:
                        try:
                            combined_text.append(extract_text_from_supported_file(inner_path, inner_ext))
                        except Exception:
                            continue
        return '\n'.join(filter(None, combined_text))
    return ''

def score_resume_against_job_description(resume_text, job_desc):
    resume_words = set(normalize_ats_text(resume_text))
    job_words = normalize_ats_text(job_desc)
    job_keywords = list(dict.fromkeys(job_words))
    if not job_keywords:
        return 0, [], []
    matched = [word for word in job_keywords if word in resume_words]
    missing = [word for word in job_keywords if word not in resume_words]
    score = round((len(matched) / len(job_keywords)) * 100)
    return score, matched, missing

# ─────────────────────────────────────────
#  PAGES (MAIN)
# ─────────────────────────────────────────
@app.route('/')
@app.route('/index.html')
@app.route('/templates/index.html')
def home():
    return render_template('index.html')

@app.route('/converter')
@app.route('/converter.html')
@app.route('/templates/converter.html')
def converter():
    return render_template('converter.html')

@app.route('/resume-ats')
@app.route('/resume-ats.html')
@app.route('/templates/resume-ats.html')
def resume_ats():
    return render_template('resume-ats.html')

@app.route('/hr-helper')
@app.route('/hr-helper.html')
@app.route('/templates/hr-helper.html')
def hr_helper():
    return render_template('hr-helper.html')

# ─────────────────────────────────────────
#  NEW PAGES - PRIVACY POLICY, TERMS, ETC.
# ─────────────────────────────────────────

@app.route('/privacy-policy')
@app.route('/privacy-policy.html')
@app.route('/templates/privacy-policy.html')
def privacy_policy():
    return render_template('privacy-policy.html')

@app.route('/terms')
@app.route('/terms.html')
@app.route('/templates/terms.html')
def terms():
    return render_template('terms.html')

@app.route('/contact')
@app.route('/contact.html')
@app.route('/templates/contact.html')
def contact():
    return render_template('contact.html')

@app.route('/about')
@app.route('/about.html')
@app.route('/templates/about.html')
def about():
    return render_template('about.html')

@app.route('/blogs')
@app.route('/blogs.html')
@app.route('/templates/blogs.html')
def blogs():
    return render_template('blogs.html')

@app.route('/blogs/<slug>')
@app.route('/blogs/<slug>.html')
def blog_post(slug):
    if slug.endswith('.html'):
        slug = slug[:-5]
    template_name = BLOG_POST_TEMPLATES.get(slug)
    if not template_name:
        abort(404)
    return render_template(template_name)

@app.route('/disclaimer')
@app.route('/disclaimer.html')
@app.route('/templates/disclaimer.html')
def disclaimer():
    return render_template('disclaimer.html')


@app.route('/templates/<path:template_file>')
def templates_alias(template_file):
    template_name = template_file.strip().lower()
    if template_name in PAGE_TEMPLATES.values():
        return render_template(template_name)
    abort(404)


@app.route('/<page>.html')
def html_alias(page):
    template_name = PAGE_TEMPLATES.get(page.strip().lower())
    if template_name:
        return render_template(template_name)
    abort(404)

@app.route('/sitemap.xml')
def sitemap():
    return send_from_directory(app.root_path, 'sitemap.xml', mimetype='application/xml')

@app.route('/robots.txt')
def robots():
    return send_from_directory(app.root_path, 'robots.txt', mimetype='text/plain')

@app.route('/ads.txt')
def ads_txt():
    return send_from_directory(app.root_path, 'ads.txt', mimetype='text/plain')

@app.route('/tracking-consent.js')
def tracking_consent_js():
    return send_from_directory(app.root_path, 'tracking-consent.js', mimetype='application/javascript')

@app.route('/health')
def health():
    return jsonify({'status': 'ok'})

@app.route('/404')
def custom_404_page():
    return send_from_directory(app.root_path, '404.html')

@app.errorhandler(404)
def not_found(_error):
    return send_from_directory(app.root_path, '404.html'), 404

@app.route('/favicon.svg')
def favicon_svg():
    return send_from_directory(app.root_path, 'favicon.svg', mimetype='image/svg+xml')

@app.route('/favicon.ico')
def favicon_ico():
    return send_from_directory(app.root_path, 'favicon.svg', mimetype='image/svg+xml')

# ─────────────────────────────────────────
#  OLD ROUTES REDIRECT
# ─────────────────────────────────────────

@app.route('/resume')
def resume_old():
    return redirect('/resume-ats')

@app.route('/ats')
def ats_old():
    return redirect('/hr-helper')

# ─────────────────────────────────────────
#  RESUME BUILDER - DOCX EXPORT
# ─────────────────────────────────────────
@app.route('/export-resume-docx', methods=['POST'])
def export_resume_docx():
    try:
        docx = importlib.import_module('docx')
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        data = request.json
        
        doc = docx.Document()
        
        # Title
        title = doc.add_heading(data.get('name', 'Resume'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact
        contact_parts = []
        if data.get('email'):
            contact_parts.append(data.get('email'))
        if data.get('phone'):
            contact_parts.append(data.get('phone'))
        if data.get('location'):
            contact_parts.append(data.get('location'))
        if data.get('linkedin'):
            contact_parts.append(data.get('linkedin'))
        
        if contact_parts:
            contact = doc.add_paragraph(' | '.join(contact_parts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Summary
        if data.get('summary'):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(data.get('summary'))
        
        # Education
        if data.get('education'):
            doc.add_heading('Education', level=1)
            for edu in data.get('education'):
                if edu.get('degree') or edu.get('school'):
                    p = doc.add_paragraph()
                    degree_text = edu.get('degree', '')
                    school_text = edu.get('school', '')
                    if degree_text and school_text:
                        p.add_run(f"{degree_text} - {school_text}").bold = True
                    elif degree_text:
                        p.add_run(degree_text).bold = True
                    elif school_text:
                        p.add_run(school_text).bold = True
                    
                    year_text = edu.get('year', '')
                    gpa_text = edu.get('gpa', '')
                    if year_text or gpa_text:
                        info = []
                        if year_text:
                            info.append(year_text)
                        if gpa_text:
                            info.append(f"GPA: {gpa_text}")
                        doc.add_paragraph(' | '.join(info))
                    doc.add_paragraph()
        
        # Experience
        if data.get('experience'):
            doc.add_heading('Work Experience', level=1)
            for exp in data.get('experience'):
                if exp.get('title') or exp.get('company'):
                    p = doc.add_paragraph()
                    title_text = exp.get('title', '')
                    company_text = exp.get('company', '')
                    if title_text and company_text:
                        p.add_run(f"{title_text} at {company_text}").bold = True
                    elif title_text:
                        p.add_run(title_text).bold = True
                    elif company_text:
                        p.add_run(company_text).bold = True
                    
                    if exp.get('duration'):
                        doc.add_paragraph(exp.get('duration'))
                    
                    if exp.get('description'):
                        doc.add_paragraph(exp.get('description'), style='List Bullet')
                    doc.add_paragraph()
        
        # Skills
        if data.get('skills'):
            doc.add_heading('Skills', level=1)
            skills_list = [s.strip() for s in data.get('skills').split(',') if s.strip()]
            doc.add_paragraph(', '.join(skills_list))
        
        # Custom Sections
        if data.get('customSections'):
            for section in data.get('customSections'):
                if section.get('title') and section.get('items'):
                    doc.add_heading(section.get('title'), level=1)
                    for item in section.get('items'):
                        if item.get('text'):
                            doc.add_paragraph(item.get('text'), style='List Bullet')
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name='resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ─────────────────────────────────────────
#  OFFICE CONVERTERS
# ─────────────────────────────────────────
@app.route('/word-to-pdf', methods=['POST'])
def word_to_pdf():
    if not is_windows_runtime():
        return windows_only_response('Word to PDF')
    try:
        docx2pdf = importlib.import_module('docx2pdf')
        docx_to_pdf = docx2pdf.convert
        p = save(request.files['file'], 'input.docx')
        docx_to_pdf(p, os.path.join(UPLOAD,'output.pdf'))
        return send_file(os.path.join(UPLOAD,'output.pdf'), as_attachment=True, download_name='converted.pdf')
    except Exception as e:
        return jsonify({'error': f'Word to PDF failed: {str(e)}'}), 500

@app.route('/pdf-to-word', methods=['POST'])
def pdf_to_word():
    try:
        pdf2docx = importlib.import_module('pdf2docx')
        PDFToWordConverter = pdf2docx.Converter
        p = save(request.files['file'], 'input.pdf')
        cv = PDFToWordConverter(p)
        out = os.path.join(UPLOAD,'output.docx')
        cv.convert(out); cv.close()
        return send_file(out, as_attachment=True, download_name='converted.docx')
    except Exception as e:
        return jsonify({'error': f'PDF to Word failed: {str(e)}'}), 500

@app.route('/ppt-to-pdf', methods=['POST'])
def ppt_to_pdf():
    if not is_windows_runtime():
        return windows_only_response('PPT to PDF')
    try:
        p = save(request.files['file'], 'input.pptx')
        comtypes_client = importlib.import_module('comtypes.client')
        import sys
        sys.modules.setdefault('comtypes.client', comtypes_client)
        ppt = comtypes_client.CreateObject("Powerpoint.Application")
        ppt.Visible = 1
        deck = ppt.Presentations.Open(os.path.abspath(p))
        out  = os.path.abspath(os.path.join(UPLOAD,'output.pdf'))
        deck.SaveAs(out, 32); deck.Close(); ppt.Quit()
        return send_file(out, as_attachment=True, download_name='converted.pdf')
    except Exception as e:
        return jsonify({'error': f'PPT to PDF failed: {str(e)}'}), 500

@app.route('/excel-to-pdf', methods=['POST'])
def excel_to_pdf():
    if not is_windows_runtime():
        return windows_only_response('Excel to PDF')
    try:
        p = save(request.files['file'], 'input.xlsx')
        comtypes_client = importlib.import_module('comtypes.client')
        import sys
        sys.modules.setdefault('comtypes.client', comtypes_client)
        xl  = comtypes_client.CreateObject("Excel.Application")
        xl.Visible = 0
        wb  = xl.Workbooks.Open(os.path.abspath(p))
        out = os.path.abspath(os.path.join(UPLOAD,'output.pdf'))
        wb.ExportAsFixedFormat(0, out); wb.Close(False); xl.Quit()
        return send_file(out, as_attachment=True, download_name='converted.pdf')
    except Exception as e:
        return jsonify({'error': f'Excel to PDF failed: {str(e)}'}), 500

@app.route('/rtf-to-docx', methods=['POST'])
def rtf_to_docx():
    if not is_windows_runtime():
        return windows_only_response('RTF to DOCX')
    try:
        p = save(request.files['file'], 'input.rtf')
        comtypes_client = importlib.import_module('comtypes.client')
        import sys
        sys.modules.setdefault('comtypes.client', comtypes_client)
        word = comtypes_client.CreateObject("Word.Application")
        word.Visible = 0
        doc  = word.Documents.Open(os.path.abspath(p))
        out  = os.path.abspath(os.path.join(UPLOAD,'output.docx'))
        doc.SaveAs2(out, 16); doc.Close(); word.Quit()
        return send_file(out, as_attachment=True, download_name='converted.docx')
    except Exception as e:
        return jsonify({'error': f'RTF to DOCX failed: {str(e)}'}), 500

@app.route('/docx-to-txt', methods=['POST'])
def docx_to_txt():
    try:
        p = save(request.files['file'], 'input.docx')
        text = extract_text_from_docx(p)
        out = os.path.join(UPLOAD, 'output.txt')
        with open(out, 'w', encoding='utf-8') as f:
            f.write(text)
        return send_file(out, as_attachment=True, download_name='converted.txt', mimetype='text/plain')
    except Exception as e:
        return jsonify({'error': f'DOCX to TXT failed: {str(e)}'}), 500

@app.route('/excel-to-csv', methods=['POST'])
def excel_to_csv():
    try:
        pd = importlib.import_module('pandas')
        p = save(request.files['file'], 'input.xlsx')
        df = pd.read_excel(p)
        out = os.path.join(UPLOAD, 'output.csv')
        df.to_csv(out, index=False, encoding='utf-8-sig')
        return send_file(out, as_attachment=True, download_name='converted.csv', mimetype='text/csv')
    except Exception as e:
        return jsonify({'error': f'Excel to CSV failed: {str(e)}'}), 500

@app.route('/word-to-html', methods=['POST'])
def word_to_html():
    try:
        p = save(request.files['file'], 'input.docx')
        text = extract_text_from_docx(p)
        out = os.path.join(UPLOAD, 'converted.html')
        lines = [f"<p>{html.escape(line.strip())}</p>" for line in text.split('\n') if line.strip()]
        html_content = "<!DOCTYPE html><html><head><meta charset='UTF-8'><title>Converted Document</title></head><body>" + ''.join(lines) + "</body></html>"
        with open(out, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return send_file(out, as_attachment=True, download_name='converted.html', mimetype='text/html')
    except Exception as e:
        return jsonify({'error': f'Word to HTML failed: {str(e)}'}), 500

@app.route('/csv-to-json', methods=['POST'])
def csv_to_json():
    try:
        p = save(request.files['file'], 'input.csv')
        out = os.path.join(UPLOAD, 'converted.json')
        with open(p, 'r', encoding='utf-8', errors='ignore', newline='') as f:
            reader = csv.DictReader(f)
            data = list(reader)
        with open(out, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return send_file(out, as_attachment=True, download_name='converted.json', mimetype='application/json')
    except Exception as e:
        return jsonify({'error': f'CSV to JSON failed: {str(e)}'}), 500

@app.route('/json-to-csv', methods=['POST'])
def json_to_csv():
    try:
        p = save(request.files['file'], 'input.json')
        out = os.path.join(UPLOAD, 'converted.csv')
        with open(p, 'r', encoding='utf-8', errors='ignore') as f:
            data = json.load(f)

        if isinstance(data, dict):
            rows = [data]
        elif isinstance(data, list):
            rows = data
        else:
            rows = [{'value': data}]

        if rows and isinstance(rows[0], dict):
            fieldnames = sorted({k for row in rows if isinstance(row, dict) for k in row.keys()})
            with open(out, 'w', encoding='utf-8', newline='') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                for row in rows:
                    writer.writerow(row if isinstance(row, dict) else {'value': row})
        else:
            with open(out, 'w', encoding='utf-8', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(['value'])
                for row in rows:
                    writer.writerow([row])

        return send_file(out, as_attachment=True, download_name='converted.csv', mimetype='text/csv')
    except Exception as e:
        return jsonify({'error': f'JSON to CSV failed: {str(e)}'}), 500

@app.route('/json-formatter', methods=['POST'])
def json_formatter():
    try:
        p = save(request.files['file'], 'input.json')
        mode = request.form.get('mode', 'pretty')
        out = os.path.join(UPLOAD, 'formatted.json')
        with open(p, 'r', encoding='utf-8', errors='ignore') as f:
            payload = json.load(f)

        formatted = json.dumps(payload, ensure_ascii=False, indent=(None if mode == 'minify' else 2))
        with open(out, 'w', encoding='utf-8') as f:
            f.write(formatted)
        return send_file(out, as_attachment=True, download_name='formatted.json', mimetype='application/json')
    except Exception as e:
        return jsonify({'error': f'JSON Formatter failed: {str(e)}'}), 500

# ─────────────────────────────────────────
#  PDF TOOLS
# ─────────────────────────────────────────
@app.route('/merge-pdf', methods=['POST'])
def merge_pdf():
    PyPDF2 = importlib.import_module('PyPDF2')
    files  = request.files.getlist('files')
    writer = PyPDF2.PdfWriter()
    for f in files:
        p = save(f, f.filename)
        reader = PyPDF2.PdfReader(p)
        for page in reader.pages:
            writer.add_page(page)
    out = os.path.join(UPLOAD,'merged.pdf')
    with open(out,'wb') as o:
        writer.write(o)
    return send_file(out, as_attachment=True, download_name='merged.pdf')

@app.route('/split-pdf', methods=['POST'])
def split_pdf():
    PyPDF2 = importlib.import_module('PyPDF2')
    p      = save(request.files['file'], 'input.pdf')
    reader = PyPDF2.PdfReader(p)
    zpath  = os.path.join(UPLOAD,'split_pages.zip')
    with zipfile.ZipFile(zpath,'w') as zf:
        for i, page in enumerate(reader.pages):
            writer = PyPDF2.PdfWriter()
            writer.add_page(page)
            ppath = os.path.join(UPLOAD, f'page_{i+1}.pdf')
            with open(ppath,'wb') as f:
                writer.write(f)
            zf.write(ppath, f'page_{i+1}.pdf')
    return send_file(zpath, as_attachment=True, download_name='split_pages.zip')

@app.route('/compress-pdf', methods=['POST'])
def compress_pdf():
    PyPDF2 = importlib.import_module('PyPDF2')
    p      = save(request.files['file'], 'input.pdf')
    reader = PyPDF2.PdfReader(p)
    writer = PyPDF2.PdfWriter()
    for page in reader.pages:
        page.compress_content_streams()
        writer.add_page(page)
    out = os.path.join(UPLOAD,'compressed.pdf')
    with open(out,'wb') as f:
        writer.write(f)
    return send_file(out, as_attachment=True, download_name='compressed.pdf')

@app.route('/rotate-pdf', methods=['POST'])
def rotate_pdf():
    PyPDF2 = importlib.import_module('PyPDF2')
    p       = save(request.files['file'], 'input.pdf')
    degrees = int(request.form.get('degrees', 90))
    reader  = PyPDF2.PdfReader(p)
    writer  = PyPDF2.PdfWriter()
    for page in reader.pages:
        page.rotate(degrees)
        writer.add_page(page)
    out = os.path.join(UPLOAD,'rotated.pdf')
    with open(out,'wb') as f:
        writer.write(f)
    return send_file(out, as_attachment=True, download_name='rotated.pdf')

@app.route('/lock-pdf', methods=['POST'])
def lock_pdf():
    PyPDF2 = importlib.import_module('PyPDF2')
    p        = save(request.files['file'], 'input.pdf')
    password = request.form.get('password', '1234')
    reader   = PyPDF2.PdfReader(p)
    writer   = PyPDF2.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    out = os.path.join(UPLOAD,'locked.pdf')
    with open(out,'wb') as f:
        writer.write(f)
    return send_file(out, as_attachment=True, download_name='locked.pdf')

@app.route('/unlock-pdf', methods=['POST'])
def unlock_pdf():
    PyPDF2 = importlib.import_module('PyPDF2')
    p        = save(request.files['file'], 'input.pdf')
    password = request.form.get('password', '')
    reader   = PyPDF2.PdfReader(p)
    if reader.is_encrypted:
        reader.decrypt(password)
    writer = PyPDF2.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    out = os.path.join(UPLOAD,'unlocked.pdf')
    with open(out,'wb') as f:
        writer.write(f)
    return send_file(out, as_attachment=True, download_name='unlocked.pdf')

@app.route('/watermark-pdf', methods=['POST'])
def watermark_pdf():
    PyPDF2 = importlib.import_module('PyPDF2')
    pagesizes = importlib.import_module('reportlab.lib.pagesizes')
    colors = importlib.import_module('reportlab.lib.colors')
    platypus = importlib.import_module('reportlab.platypus')
    styles = importlib.import_module('reportlab.lib.styles')
    p    = save(request.files['file'], 'input.pdf')
    text = request.form.get('watermark', 'CONFIDENTIAL')
    reader = PyPDF2.PdfReader(p)
    writer = PyPDF2.PdfWriter()
    wm_path = os.path.join(UPLOAD,'wm.pdf')
    c = platypus.SimpleDocTemplate(wm_path, pagesize=pagesizes.A4)
    style = styles.ParagraphStyle('wm', fontSize=40, textColor=colors.Color(0.7,0.7,0.7,0.4),
        fontName='Helvetica-Bold')
    c.build([platypus.Paragraph(text, style)])
    wm_reader = PyPDF2.PdfReader(wm_path)
    wm_page   = wm_reader.pages[0]
    for page in reader.pages:
        page.merge_page(wm_page)
        writer.add_page(page)
    out = os.path.join(UPLOAD,'watermarked.pdf')
    with open(out,'wb') as f:
        writer.write(f)
    return send_file(out, as_attachment=True, download_name='watermarked.pdf')

@app.route('/delete-pages', methods=['POST'])
def delete_pages():
    PyPDF2 = importlib.import_module('PyPDF2')
    p     = save(request.files['file'], 'input.pdf')
    pages = request.form.get('pages', '')
    to_delete = set()
    for part in pages.split(','):
        part = part.strip()
        if '-' in part:
            a,b = part.split('-')
            to_delete.update(range(int(a)-1, int(b)))
        elif part:
            to_delete.add(int(part)-1)
    reader = PyPDF2.PdfReader(p)
    writer = PyPDF2.PdfWriter()
    for i, page in enumerate(reader.pages):
        if i not in to_delete:
            writer.add_page(page)
    out = os.path.join(UPLOAD,'deleted.pdf')
    with open(out,'wb') as f:
        writer.write(f)
    return send_file(out, as_attachment=True, download_name='deleted_pages.pdf')

@app.route('/extract-pages', methods=['POST'])
def extract_pages():
    PyPDF2 = importlib.import_module('PyPDF2')
    p     = save(request.files['file'], 'input.pdf')
    pages = request.form.get('pages', '')
    to_keep = set()
    for part in pages.split(','):
        part = part.strip()
        if '-' in part:
            a,b = part.split('-')
            to_keep.update(range(int(a)-1, int(b)))
        elif part:
            to_keep.add(int(part)-1)
    reader = PyPDF2.PdfReader(p)
    writer = PyPDF2.PdfWriter()
    for i, page in enumerate(reader.pages):
        if i in to_keep:
            writer.add_page(page)
    out = os.path.join(UPLOAD,'extracted.pdf')
    with open(out,'wb') as f:
        writer.write(f)
    return send_file(out, as_attachment=True, download_name='extracted_pages.pdf')

# ─────────────────────────────────────────
#  PDF TO IMAGE - Using PyMuPDF (No system dependencies)
# ─────────────────────────────────────────
@app.route('/pdf-to-image', methods=['POST'])
def pdf_to_image():
    try:
        fitz = importlib.import_module('fitz')
        p = save(request.files['file'], 'input.pdf')
        doc = fitz.open(p)
        zpath = os.path.join(UPLOAD, 'pdf_images.zip')
        with zipfile.ZipFile(zpath, 'w') as zf:
            for i, page in enumerate(doc):
                pix = page.get_pixmap()
                ipath = os.path.join(UPLOAD, f'page_{i+1}.jpg')
                pix.save(ipath)
                zf.write(ipath, f'page_{i+1}.jpg')
        doc.close()
        return send_file(zpath, as_attachment=True, download_name='pdf_images.zip')
    except Exception as e:
        return jsonify({'error': f'PDF to Image failed: {str(e)}. Try: pip install PyMuPDF'}), 500

@app.route('/pdf-to-txt', methods=['POST'])
def pdf_to_txt():
    p    = save(request.files['file'], 'input.pdf')
    text = extract_text_from_pdf(p)
    out  = os.path.join(UPLOAD,'output.txt')
    with open(out,'w', encoding='utf-8') as f:
        f.write(text)
    return send_file(out, as_attachment=True, download_name='converted.txt')

@app.route('/txt-to-pdf', methods=['POST'])
def txt_to_pdf():
    pagesizes = importlib.import_module('reportlab.lib.pagesizes')
    platypus = importlib.import_module('reportlab.platypus')
    styles = importlib.import_module('reportlab.lib.styles')
    units = importlib.import_module('reportlab.lib.units')
    p = save(request.files['file'], 'input.txt')
    with open(p,'r', encoding='utf-8') as f:
        text = f.read()
    out = os.path.join(UPLOAD,'output.pdf')
    doc = platypus.SimpleDocTemplate(out, pagesize=pagesizes.A4,
        rightMargin=2*units.cm, leftMargin=2*units.cm, topMargin=2*units.cm, bottomMargin=2*units.cm)
    style = styles.ParagraphStyle('body', fontSize=11, leading=16)
    story = [platypus.Paragraph(line or ' ', style) for line in text.split('\n')]
    doc.build(story)
    return send_file(out, as_attachment=True, download_name='converted.pdf')

@app.route('/add-page-numbers', methods=['POST'])
def add_page_numbers():
    PyPDF2 = importlib.import_module('PyPDF2')
    pagesizes = importlib.import_module('reportlab.lib.pagesizes')
    platypus = importlib.import_module('reportlab.platypus')
    styles = importlib.import_module('reportlab.lib.styles')
    units = importlib.import_module('reportlab.lib.units')
    colors = importlib.import_module('reportlab.lib.colors')
    p      = save(request.files['file'], 'input.pdf')
    reader = PyPDF2.PdfReader(p)
    writer = PyPDF2.PdfWriter()
    total  = len(reader.pages)
    for i, page in enumerate(reader.pages):
        pn_path = os.path.join(UPLOAD, f'pn_{i}.pdf')
        pn_doc  = platypus.SimpleDocTemplate(pn_path, pagesize=pagesizes.A4,
            rightMargin=1*units.cm, leftMargin=1*units.cm, topMargin=1*units.cm, bottomMargin=1*units.cm)
        style = styles.ParagraphStyle('pn', fontSize=9, textColor=colors.grey)
        pn_doc.build([platypus.Spacer(1, 26*units.cm), platypus.Paragraph(f'Page {i+1} of {total}', style)])
        pn_reader = PyPDF2.PdfReader(pn_path)
        page.merge_page(pn_reader.pages[0])
        writer.add_page(page)
    out = os.path.join(UPLOAD,'numbered.pdf')
    with open(out,'wb') as f:
        writer.write(f)
    return send_file(out, as_attachment=True, download_name='numbered.pdf')

# ─────────────────────────────────────────
#  IMAGE TOOLS
# ─────────────────────────────────────────
@app.route('/image-to-pdf', methods=['POST'])
def image_to_pdf():
    p   = save(request.files['file'], 'input_img.png')
    img = Image.open(p).convert('RGB')
    out = os.path.join(UPLOAD,'output.pdf')
    img.save(out,'PDF')
    return send_file(out, as_attachment=True, download_name='converted.pdf')

# ─────────────────────────────────────────
#  IMAGE TO WORD - Using EasyOCR (No Tesseract needed)
# ─────────────────────────────────────────
@app.route('/image-to-word', methods=['POST'])
def image_to_word():
    try:
        easyocr = importlib.import_module('easyocr')
        docx = importlib.import_module('docx')
        p = save(request.files['file'], 'input_img.png')
        reader = easyocr.Reader(['en'])
        result = reader.readtext(p)
        text = ' '.join([item[1] for item in result])
        
        d = docx.Document()
        d.add_heading('Extracted Text', 0)
        for line in text.split('\n'):
            if line.strip():
                d.add_paragraph(line.strip())
        out = os.path.join(UPLOAD, 'output.docx')
        d.save(out)
        return send_file(out, as_attachment=True, download_name='converted.docx')
    except Exception as e:
        return jsonify({'error': f'OCR to Word failed: {str(e)}. Try: pip install easyocr'}), 500

@app.route('/image-to-excel', methods=['POST'])
def image_to_excel():
    try:
        easyocr = importlib.import_module('easyocr')
        pd = importlib.import_module('pandas')
        p = save(request.files['file'], 'input_img.png')
        reader = easyocr.Reader(['en'])
        result = reader.readtext(p)
        text = ' '.join([item[1] for item in result])
        
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        rows = [re.split(r'\t|  {2,}', l) for l in lines]
        out = os.path.join(UPLOAD, 'output.xlsx')
        pd.DataFrame(rows).to_excel(out, index=False, header=False)
        return send_file(out, as_attachment=True, download_name='converted.xlsx')
    except Exception as e:
        return jsonify({'error': f'OCR to Excel failed: {str(e)}. Try: pip install easyocr'}), 500

@app.route('/ocr-to-text', methods=['POST'])
def ocr_to_text():
    try:
        easyocr = importlib.import_module('easyocr')
        p = save(request.files['file'], 'input_img.png')
        reader = easyocr.Reader(['en'])
        result = reader.readtext(p)
        text = ' '.join([item[1] for item in result])
        
        out = os.path.join(UPLOAD, 'ocr_output.txt')
        with open(out, 'w', encoding='utf-8') as f:
            f.write(text)
        return send_file(out, as_attachment=True, download_name='ocr_text.txt')
    except Exception as e:
        return jsonify({'error': f'OCR to Text failed: {str(e)}. Try: pip install easyocr'}), 500

@app.route('/jpg-to-png', methods=['POST'])
def jpg_to_png():
    p   = save(request.files['file'], 'input.jpg')
    out = os.path.join(UPLOAD,'output.png')
    Image.open(p).save(out,'PNG')
    return send_file(out, as_attachment=True, download_name='converted.png')

@app.route('/png-to-jpg', methods=['POST'])
def png_to_jpg():
    p   = save(request.files['file'], 'input.png')
    out = os.path.join(UPLOAD,'output.jpg')
    Image.open(p).convert('RGB').save(out,'JPEG')
    return send_file(out, as_attachment=True, download_name='converted.jpg')

@app.route('/webp-to-png', methods=['POST'])
def webp_to_png():
    p   = save(request.files['file'], 'input.webp')
    out = os.path.join(UPLOAD,'output.png')
    Image.open(p).save(out,'PNG')
    return send_file(out, as_attachment=True, download_name='converted.png')

@app.route('/svg-to-png', methods=['POST'])
def svg_to_png():
    try:
        cairosvg = importlib.import_module('cairosvg')
        p   = save(request.files['file'], 'input.svg')
        out = os.path.join(UPLOAD,'output.png')
        cairosvg.svg2png(url=p, write_to=out)
        return send_file(out, as_attachment=True, download_name='converted.png')
    except Exception as e:
        return jsonify({'error': f'SVG to PNG failed: {str(e)}. Try: pip install cairosvg'}), 500

@app.route('/compress-image', methods=['POST'])
def compress_image():
    p       = save(request.files['file'], 'input_img.png')
    quality = int(request.form.get('quality', 60))
    img     = Image.open(p).convert('RGB')
    out     = os.path.join(UPLOAD,'compressed.jpg')
    img.save(out,'JPEG', quality=quality)
    return send_file(out, as_attachment=True, download_name='compressed.jpg')

@app.route('/flip-image', methods=['POST'])
def flip_image():
    try:
        p = save(request.files['file'], 'input_flip_image')
        direction = request.form.get('direction', 'horizontal')
        img = Image.open(p)
        out = os.path.join(UPLOAD, 'flipped.png')
        if direction == 'vertical':
            img = img.transpose(Image.FLIP_TOP_BOTTOM)
        else:
            img = img.transpose(Image.FLIP_LEFT_RIGHT)
        img.save(out, 'PNG')
        return send_file(out, as_attachment=True, download_name='flipped.png')
    except Exception as e:
        return jsonify({'error': f'Flip Image failed: {str(e)}'}), 500

@app.route('/download-resume/<filename>')
def download_resume(filename):
    path = os.path.join(UPLOAD, filename)
    return send_file(path, as_attachment=True, download_name=filename)

@app.route('/ats-score', methods=['POST'])
def ats_score():
    file = request.files.get('file') or request.files.get('resume')
    job_desc = request.form.get('job_desc', '') or request.form.get('jobDesc', '')

    if not file:
        return jsonify({'error': 'No file uploaded'}), 400

    filename = file.filename or 'resume'
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    path = save(file, filename)

    try:
        resume_text = extract_text_from_supported_file(path, ext)
        score, matched, missing = score_resume_against_job_description(resume_text, job_desc)
        return jsonify({'score': score, 'matched': matched, 'missing': missing, 'filename': filename})
    except Exception as e:
        return jsonify({'error': str(e), 'filename': filename}), 500

# ─────────────────────────────────────────
#  OLD RESUME MAKER (kept for compatibility)
# ─────────────────────────────────────────
@app.route('/make-resume', methods=['POST'])
def make_resume():
    pagesizes = importlib.import_module('reportlab.lib.pagesizes')
    colors = importlib.import_module('reportlab.lib.colors')
    platypus = importlib.import_module('reportlab.platypus')
    styles = importlib.import_module('reportlab.lib.styles')
    units = importlib.import_module('reportlab.lib.units')
    data = request.json
    out  = os.path.join(UPLOAD,'my_resume.pdf')
    doc  = platypus.SimpleDocTemplate(out, pagesize=pagesizes.A4,
        rightMargin=2*units.cm, leftMargin=2*units.cm, topMargin=2*units.cm, bottomMargin=2*units.cm)
    name_s    = styles.ParagraphStyle('n', fontSize=24, fontName='Helvetica-Bold', textColor=colors.HexColor('#2d3748'), spaceAfter=4)
    contact_s = styles.ParagraphStyle('c', fontSize=10, fontName='Helvetica',     textColor=colors.HexColor('#718096'), spaceAfter=2)
    section_s = styles.ParagraphStyle('s', fontSize=13, fontName='Helvetica-Bold', textColor=colors.HexColor('#2b6cb0'), spaceBefore=14, spaceAfter=4)
    body_s    = styles.ParagraphStyle('b', fontSize=10, fontName='Helvetica',     textColor=colors.HexColor('#2d3748'), spaceAfter=3, leading=15)
    bold_s    = styles.ParagraphStyle('bl',fontSize=10, fontName='Helvetica-Bold', textColor=colors.HexColor('#2d3748'), spaceAfter=1)
    hr = lambda: platypus.HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0'), spaceAfter=5)
    story = []
    story.append(platypus.Paragraph(data.get('name',''), name_s))
    story.append(platypus.Paragraph(' | '.join(filter(None,[data.get('email',''),data.get('phone',''),data.get('location',''),data.get('linkedin','')])), contact_s))
    story.append(platypus.HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2b6cb0'), spaceAfter=8))
    if data.get('summary'):
        story += [platypus.Paragraph('PROFESSIONAL SUMMARY',section_s), hr(), platypus.Paragraph(data['summary'],body_s)]
    if data.get('education'):
        story += [platypus.Paragraph('EDUCATION',section_s), hr()]
        for e in data['education']:
            story += [platypus.Paragraph(f"{e.get('degree','')} — {e.get('school','')}",bold_s), platypus.Paragraph(f"{e.get('year','')}  {e.get('gpa','')}",body_s)]
    if data.get('experience'):
        story += [platypus.Paragraph('WORK EXPERIENCE',section_s), hr()]
        for e in data['experience']:
            story += [platypus.Paragraph(f"{e.get('title','')} — {e.get('company','')}",bold_s), platypus.Paragraph(e.get('duration',''),body_s), platypus.Paragraph(e.get('description',''),body_s), platypus.Spacer(1,4)]
    if data.get('skills'):
        story += [platypus.Paragraph('SKILLS',section_s), hr(), platypus.Paragraph(data['skills'],body_s)]
    if data.get('projects'):
        story += [platypus.Paragraph('PROJECTS',section_s), hr()]
        for p in data['projects']:
            story += [platypus.Paragraph(p.get('title',''),bold_s), platypus.Paragraph(p.get('description',''),body_s), platypus.Spacer(1,4)]
    doc.build(story)
    return send_file(out, as_attachment=True, download_name='my_resume.pdf')

# ─────────────────────────────────────────
#  LINKEDIN - File Upload API
# ─────────────────────────────────────────
@app.route('/upload-linkedin-file', methods=['POST'])
def upload_linkedin_file():
    try:
        file = request.files['file']
        if not file:
            return jsonify({'error': 'No file uploaded'}), 400
        
        filename = file.filename
        ext = filename.rsplit('.', 1)[-1].lower()
        path = save(file, filename)

        text = extract_text_from_supported_file(path, ext)
        if not text.strip():
            return jsonify({'error': 'Could not extract text from the uploaded file'}), 400

        if request.args.get('transport') == 'message':
            token = request.form.get('token', '')
            payload = json.dumps({
                'type': 'ats-upload-result',
                'token': token,
                'filename': filename,
                'text': text,
                'error': None,
            })
            html = f"""<!doctype html>
<html>
<body>
<script>
window.parent.postMessage({payload}, '*');
</script>
</body>
</html>"""
            response = make_response(html)
            response.headers['Content-Type'] = 'text/html; charset=utf-8'
            return response
        
        return jsonify({'text': text})
    except Exception as e:
        if request.args.get('transport') == 'message':
            token = request.form.get('token', '')
            payload = json.dumps({
                'type': 'ats-upload-result',
                'token': token,
                'filename': request.files['file'].filename if 'file' in request.files else '',
                'text': '',
                'error': str(e),
            })
            html = f"""<!doctype html>
<html>
<body>
<script>
window.parent.postMessage({payload}, '*');
</script>
</body>
</html>"""
            response = make_response(html)
            response.headers['Content-Type'] = 'text/html; charset=utf-8'
            return response
        return jsonify({'error': str(e)}), 500

# ─────────────────────────────────────────
#  RUN SERVER ON PORT 5001
# ─────────────────────────────────────────
if __name__ == '__main__':
    app.run(debug=True, port=5001)
